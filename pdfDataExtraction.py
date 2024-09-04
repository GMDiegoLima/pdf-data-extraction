import os
import io
import fitz
import logging
import pytesseract
from docx import Document
from docx.shared import Pt
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', datefmt='%H:%M:%S')
pdfs_folder:str = 'PDFs to analyse'
output_folder:str = 'Final results'
font_name:str = 'Times New Roman'

def create_docx():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(12)
    return doc

def ensure_folder_exists(folder_path):
    if not os.path.exists(folder_path):
        logging.info(f'Created {output_folder} folder to save the final files results')
        os.makedirs(folder_path)

def process_pdfs(pdfsfolder, output_name, output_folder, process_function, log_message):
    logging.info(f'Checking the PDFs in the {pdfsfolder} folder')
    ensure_folder_exists(output_folder)
    for index, pdf_filename in enumerate(os.listdir(pdfsfolder)):
        if pdf_filename.endswith('.pdf'):
            logging.info(f'Extracting data from {pdf_filename}, please wait...')
            process_function(output_name, output_folder, os.path.join(pdfsfolder, pdf_filename), index)
    logging.info(log_message)
    return

def pdf_images_to_text(output_filename: str,  output_folder:str, filepath: str, index: int):
    pdf = fitz.open(filepath)
    canvs = canvas.Canvas(os.path.join(output_folder, f'{output_filename}_{str(index)}.pdf'), pagesize=letter)
    canvs.setFont('Helvetica', 12)

    for page_number in range(len(pdf)):
        page = pdf.load_page(page_number)
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = pdf.extract_image(xref)
            bytes_image = base_image["image"]
            image = Image.open(io.BytesIO(bytes_image))
            text = pytesseract.image_to_string(image)

            for idx_line, line in enumerate(text.splitlines(), start=1):
                canvs.drawString(100, 750 - 15 * idx_line, line)
            canvs.showPage()

    logging.info(f'Saving {output_filename}_{str(index)}.pdf')
    canvs.save()
    return

def docx_images_to_text(output_filename: str,  output_folder:str, filepath: str, index: int):
    docx = create_docx()
    pdf = fitz.open(filepath)
    for page_number in range(len(pdf)):
        page = pdf.load_page(page_number)
        for img in page.get_images(full=True):
            xref = img[0]
            base_image = pdf.extract_image(xref)
            bytes_image = base_image["image"]
            image = Image.open(io.BytesIO(bytes_image))
            text = pytesseract.image_to_string(image)

            docx.add_paragraph(text)
            docx.add_page_break()

    docx.save(os.path.join(output_folder, f'{output_filename}.docx'))
    return

def pdf_to_docx(output_filename: str,  output_folder:str, filepath: str, index: int):
    pdf = fitz.open(filepath)
    docx = create_docx()
    for page_num in range(pdf.page_count):
        page = pdf.load_page(page_num)
        text = page.get_text()
        
        docx.add_paragraph(text)
        docx.add_page_break()
    docx.save(os.path.join(output_folder, f'{output_filename}.docx'))
    return

if __name__ == '__main__':
    user_choice:int = int(input(
        'Please select an option:\n'
        '1. Extract text from images in a PDF and save it as a new PDF\n'
        '2. Extract text from images in a PDF and save it as a DOCX file\n'
        '3. Convert the entire PDF document to a DOCX file\n'
        'Enter your choice (1-3):\n'))
    output_name:str = str(input(
        'Please, type the name of the output file WITHOUT the final extension (.pdf, .docx):\n'))

    match user_choice:
        case 1:
            process_pdfs(
                pdfs_folder,
                output_name,
                output_folder,
                pdf_images_to_text,
                f'All the images of each PDF file in {pdfs_folder} have been converted to PDFs with texts')
        case 2:
            process_pdfs(
                pdfs_folder,
                output_name,
                output_folder,
                docx_images_to_text,
                f'All the images of each PDF file in {pdfs_folder} have been converted to DOCX')
        case 3:
            process_pdfs(
                pdfs_folder,
                output_name,
                output_folder,
                pdf_to_docx,
                f'All the images of each PDF file in {pdfs_folder} have been converted to DOCX')
        case _:
            logging.error('Sorry, but you typed a invalid input, the tool only accept numbers between 1 to 3')

